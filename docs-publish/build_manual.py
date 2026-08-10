from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
PAGES = ROOT / "feishu-pages"
OUTPUT = ROOT / "照片流-用户手册.docx"

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
INK = RGBColor(0x1F, 0x29, 0x37)
MUTED = RGBColor(0x6B, 0x72, 0x80)
LIGHT = "F4F6F9"
GOLD = RGBColor(0xA1, 0x6B, 0x00)


def set_run_font(run, name="Calibri", east_asia="Microsoft YaHei", size=None, color=None,
                 bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def style_paragraph(style, font_size, color, before, after, line=1.25, bold=False):
    style.font.name = "Calibri"
    style.font.size = Pt(font_size)
    style.font.color.rgb = color
    style.font.bold = bold
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    pf = style.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def shade_paragraph(paragraph, fill):
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), fill)
    ppr.append(OxmlElement("w:keepNext"))


def set_cell_like_padding(paragraph, before=6, after=6, left=8, right=8):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.left_indent = Pt(left)
    pf.right_indent = Pt(right)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def setup_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.space_after = Pt(0)
    set_run_font(hp.add_run("照片流 PhotoFlow · 用户手册"), size=8.5, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    set_run_font(fp.add_run("第 "), size=8.5, color=MUTED)
    add_field(fp, "PAGE")
    set_run_font(fp.add_run(" 页"), size=8.5, color=MUTED)


def configure_styles(doc):
    style_paragraph(doc.styles["Normal"], 11, INK, 0, 6, 1.25)
    style_paragraph(doc.styles["Heading 1"], 16, BLUE, 18, 10, 1.1, True)
    style_paragraph(doc.styles["Heading 2"], 13, BLUE, 14, 7, 1.1, True)
    style_paragraph(doc.styles["Heading 3"], 12, DARK_BLUE, 10, 5, 1.1, True)
    for name in ("Heading 1", "Heading 2", "Heading 3"):
        doc.styles[name].paragraph_format.keep_with_next = True
        doc.styles[name].paragraph_format.keep_together = True


def new_numbering(doc, ordered=False):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids or [0]) + 1
    num_id = max(num_ids or [0]) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1." if ordered else "•")
    lvl.append(lvl_text)
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    lvl.append(jc)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    ppr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    ppr.append(spacing)
    lvl.append(ppr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    numpr = ppr.find(qn("w:numPr"))
    if numpr is None:
        numpr = OxmlElement("w:numPr")
        ppr.append(numpr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    nid = OxmlElement("w:numId")
    nid.set(qn("w:val"), str(num_id))
    numpr.append(ilvl)
    numpr.append(nid)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25


def add_inline(paragraph, text):
    pattern = re.compile(r"(\*\*.+?\*\*|`.+?`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            set_run_font(paragraph.add_run(text[pos:match.start()]), size=11, color=INK)
        token = match.group(0)
        if token.startswith("**"):
            set_run_font(paragraph.add_run(token[2:-2]), size=11, color=INK, bold=True)
        else:
            set_run_font(paragraph.add_run(token[1:-1]), name="Consolas", east_asia="Microsoft YaHei", size=9.5, color=DARK_BLUE)
        pos = match.end()
    if pos < len(text):
        set_run_font(paragraph.add_run(text[pos:]), size=11, color=INK)


def add_picture(doc, path, alt):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    shape = run.add_picture(str(path), width=Inches(6.35))
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", alt)
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(0)
    caption.paragraph_format.space_after = Pt(8)
    set_run_font(caption.add_run(f"图：{alt}"), size=9, color=MUTED)


def add_cover(doc):
    for _ in range(5):
        doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    set_run_font(kicker.add_run("PHOTOFLOW USER GUIDE"), size=10.5, color=GOLD, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    set_run_font(title.add_run("照片流 PhotoFlow"), size=30, color=DARK_BLUE, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    set_run_font(subtitle.add_run("产品介绍与完整使用教程"), size=15, color=BLUE)

    lead = doc.add_paragraph()
    lead.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lead.paragraph_format.space_after = Pt(80)
    set_run_font(lead.add_run("面向摄影师的本地项目管理与素材工作流工具"), size=11, color=MUTED, italic=True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(4)
    set_run_font(meta.add_run("适用平台：Windows 10/11 x64"), size=10.5, color=INK, bold=True)
    meta2 = doc.add_paragraph()
    meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(meta2.add_run("界面依据：当前开发版 · 截图日期 2026-08-09"), size=9.5, color=MUTED)


def add_front_matter(doc, page_files):
    doc.add_page_break()
    doc.add_heading("使用前必读", level=1)
    p = doc.add_paragraph()
    set_cell_like_padding(p, 7, 7, 9, 9)
    shade_paragraph(p, "FFF4E5")
    set_run_font(p.add_run("重要："), size=11, color=GOLD, bold=True)
    set_run_font(p.add_run("软件仍在持续开发。首次使用前请备份重要数据，并使用素材副本验证自己的导入、选片、版本和归档流程。"), size=11, color=INK)
    p2 = doc.add_paragraph()
    set_cell_like_padding(p2, 7, 7, 9, 9)
    shade_paragraph(p2, "FFF4E5")
    set_run_font(p2.add_run("文件风险："), size=11, color=GOLD, bold=True)
    set_run_font(p2.add_run("导入设置可能删除来源文件。第一次使用时建议关闭“导入后默认删除源文件”，确认目标和独立备份无误后再调整。"), size=11, color=INK)

    doc.add_heading("目录", level=1)
    num_id = new_numbering(doc, ordered=True)
    for file in page_files:
        title = file.stem.split("-", 1)[1]
        p = doc.add_paragraph()
        add_inline(p, title)
        apply_numbering(p, num_id)


def parse_markdown(doc, path):
    lines = path.read_text(encoding="utf-8").splitlines()
    list_kind = None
    list_num = None
    in_code = False
    code_lines = []
    for line in lines:
        if line.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(7)
                set_cell_like_padding(p, 6, 6, 8, 8)
                shade_paragraph(p, "F2F4F7")
                set_run_font(p.add_run("\n".join(code_lines)), name="Consolas", east_asia="Microsoft YaHei", size=9.3, color=INK)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue

        image_match = re.match(r"!\[(.*?)\]\((.*?)\)", line)
        if image_match:
            image_path = (path.parent / image_match.group(2)).resolve()
            add_picture(doc, image_path, image_match.group(1))
            list_kind = None
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            list_kind = None
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            list_kind = None
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            list_kind = None
            continue

        bullet = re.match(r"^-\s+(.+)$", line)
        number = re.match(r"^\d+\.\s+(.+)$", line)
        if bullet or number:
            kind = "number" if number else "bullet"
            if list_kind != kind:
                list_kind = kind
                list_num = new_numbering(doc, ordered=kind == "number")
            p = doc.add_paragraph()
            add_inline(p, (number or bullet).group(1))
            apply_numbering(p, list_num)
            continue

        list_kind = None
        if not line.strip():
            continue
        p = doc.add_paragraph()
        add_inline(p, line.strip())


def build():
    page_files = sorted(PAGES.glob("*.md"))
    doc = Document()
    setup_page(doc)
    configure_styles(doc)
    props = doc.core_properties
    props.title = "照片流 PhotoFlow 用户手册"
    props.subject = "产品介绍、安装与完整使用教程"
    props.author = "照片流"
    props.keywords = "照片流, PhotoFlow, 摄影, 项目管理, 使用教程"

    add_cover(doc)
    add_front_matter(doc, page_files)
    for page in page_files:
        doc.add_page_break()
        parse_markdown(doc, page)

    # Prevent accidental extra section geometry drift.
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
